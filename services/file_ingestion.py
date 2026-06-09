from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document


def _file_bytes(file) -> bytes:
    if isinstance(file, bytes):
        return file
    if hasattr(file, "getvalue"):
        return file.getvalue()
    return file.read()


def read_txt(file) -> str:
    data = _file_bytes(file)
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError:
        return data.decode("cp1252")


def read_csv(file) -> pd.DataFrame:
    data = _file_bytes(file)
    try:
        return pd.read_csv(BytesIO(data))
    except UnicodeDecodeError:
        return pd.read_csv(BytesIO(data), encoding="cp1252")


def read_xlsx(file) -> dict[str, pd.DataFrame]:
    data = _file_bytes(file)
    workbook = pd.ExcelFile(BytesIO(data), engine="openpyxl")
    return {
        sheet_name: pd.read_excel(workbook, sheet_name=sheet_name)
        for sheet_name in workbook.sheet_names
    }


def read_docx(file) -> str:
    document = Document(BytesIO(_file_bytes(file)))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text for cell in row.cells]
            if any(value.strip() for value in values):
                blocks.append(" | ".join(values))
    return "\n".join(blocks)


def normalise_dataframe_to_text(df: pd.DataFrame) -> str:
    blocks = []
    for _, row in df.iterrows():
        values = [
            f"{column}: {value}"
            for column, value in row.items()
            if pd.notna(value) and str(value).strip()
        ]
        if values:
            blocks.append(" | ".join(values))
    return "\n".join(blocks)


def _dataframe_records(df: pd.DataFrame, prefix: str = "Row") -> list[dict]:
    records = []
    for index, row in df.iterrows():
        values = [
            f"{column}: {value}"
            for column, value in row.items()
            if pd.notna(value) and str(value).strip()
        ]
        if values:
            records.append(
                {
                    "source_location": f"{prefix} {index + 2}",
                    "raw_feedback": " | ".join(values),
                }
            )
    return records


def _docx_content(data: bytes) -> dict:
    document = Document(BytesIO(data))
    records: list[dict] = []
    order = 0
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            order += 1
            records.append(
                {
                    "order": order,
                    "source_location": f"Paragraph {paragraph_number}",
                    "raw_feedback": paragraph.text,
                }
            )
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            values = [cell.text for cell in row.cells]
            if any(value.strip() for value in values):
                order += 1
                records.append(
                    {
                        "order": order,
                        "source_location": f"Table {table_number} Row {row_number}",
                        "raw_feedback": " | ".join(values),
                    }
                )
    records.sort(key=lambda item: item["order"])
    for record in records:
        record.pop("order", None)
    return {"records": records, "text": "\n".join(r["raw_feedback"] for r in records)}


def extract_content(uploaded_file) -> dict:
    filename = getattr(uploaded_file, "name", "upload.txt")
    suffix = Path(filename).suffix.lower()
    data = _file_bytes(uploaded_file)

    if suffix == ".txt":
        text = read_txt(data)
        records = [
            {"source_location": f"Line {number}", "raw_feedback": line}
            for number, line in enumerate(text.splitlines(), start=1)
            if line.strip()
        ]
    elif suffix == ".csv":
        dataframe = read_csv(data)
        records = _dataframe_records(dataframe)
        text = normalise_dataframe_to_text(dataframe)
    elif suffix == ".xlsx":
        sheets = read_xlsx(data)
        records = []
        text_parts = []
        for sheet_name, dataframe in sheets.items():
            sheet_records = _dataframe_records(dataframe, f"{sheet_name} Row")
            records.extend(sheet_records)
            sheet_text = normalise_dataframe_to_text(dataframe)
            if sheet_text:
                text_parts.append(f"[Sheet: {sheet_name}]\n{sheet_text}")
        text = "\n\n".join(text_parts)
    elif suffix == ".docx":
        content = _docx_content(data)
        records = content["records"]
        text = content["text"]
    else:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    preview = text[:5000]
    return {
        "file_type": suffix.lstrip("."),
        "text": text,
        "records": records,
        "preview": preview,
    }
