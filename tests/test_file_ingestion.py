from io import BytesIO

import pandas as pd
import pytest
from docx import Document

from services.file_ingestion import extract_content


class UploadedFile(BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def sample_uploads():
    text = UploadedFile(b"Issue one\nIssue two", "feedback.txt")
    csv = UploadedFile(b"feedback\nIssue one\nIssue two\n", "feedback.csv")

    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        pd.DataFrame({"feedback": ["Issue one", "Issue two"]}).to_excel(
            writer,
            sheet_name="Feedback",
            index=False,
        )
    excel = UploadedFile(excel_buffer.getvalue(), "feedback.xlsx")

    document = Document()
    document.add_paragraph("Issue one")
    document.add_paragraph("Issue two")
    word_buffer = BytesIO()
    document.save(word_buffer)
    word = UploadedFile(word_buffer.getvalue(), "feedback.docx")

    return [text, csv, excel, word]


@pytest.mark.parametrize("uploaded_file", sample_uploads())
def test_supported_file_types_create_feedback_records(uploaded_file):
    content = extract_content(uploaded_file)

    assert len(content["records"]) == 2
    assert content["records"][0]["raw_feedback"]
